#!/usr/bin/env python3
"""Generate a payment-record DOCX with one large title and stacked images.

When ``--images`` is given, works in manual mode (original behaviour).
When ``--images`` is omitted, reads ``invoice_errors.json`` to discover
连号发票 groups automatically, collects payment-record screenshots for each,
and generates one DOCX per group.
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

from _pathutil import INTERNAL_DIR, add_root_arg, resolve_path
from _matching_records import DEFAULT_MATCH_RECORD, image_paths, invoice_key, load_match_record


DEFAULT_OUTPUT = INTERNAL_DIR / "支付记录/xxx_17-24_支付记录.docx"
DEFAULT_TITLE = "xxx 17-24 支付记录"
DEFAULT_WIDTH_CM = 4.0
DEFAULT_ERRORS = Path("invoice_errors.json")
DEFAULT_RESULTS = Path("invoice_results_sorted.json")


def display_index(inv: dict) -> int:
    name = str(inv.get("更新后文件名") or "")
    m = re.match(r"(\d+)_", name)
    if m:
        return int(m.group(1))
    return int(inv.get("发票序号", 0)) + 1


def collect_images(images: list[Path], images_dir: Path | None, pattern: str, root: Path) -> list[Path]:
    collected = [path for path in images if path.exists()]
    if collected:
        return sorted(collected)
    if images_dir is None:
        raise RuntimeError("no images provided")
    images_dir = resolve_path(root, images_dir)
    if not images_dir.exists():
        raise RuntimeError(f"images_dir does not exist: {images_dir}")
    return sorted(path for path in images_dir.rglob(pattern) if path.is_file())


def index_text(indexes: list[int]) -> str:
    values = sorted(set(indexes))
    ranges: list[str] = []
    start = end = values[0]
    for value in values[1:]:
        if value == end + 1:
            end = value
            continue
        ranges.append(str(start) if start == end else f"{start}-{end}")
        start = end = value
    ranges.append(str(start) if start == end else f"{start}-{end}")
    return "&".join(ranges)


def is_high_unit_price(inv: dict) -> bool:
    if "辰景" in str(inv.get("购买方名称") or ""):
        return False
    for item in inv.get("项目列表", []):
        try:
            if float(item.get("单价")) > 1000:
                return True
        except (TypeError, ValueError):
            continue
    return False


def auto_collect_groups_from_record(errors_path: Path, results_path: Path, match_record: Path, root: Path) -> list[dict]:
    """Read invoice_errors.json and collect payment images from 匹配记录.json."""
    root = root.resolve()
    errors = json.loads(errors_path.read_text(encoding="utf-8"))
    results = json.loads(results_path.read_text(encoding="utf-8"))
    record = load_match_record(match_record)
    mapping = record.get("发票映射", {})
    invoices_by_source = {inv["文件名"]: inv for inv in results.get("发票信息", [])}

    groups: list[dict] = []
    candidates: list[list[dict]] = []
    for entry in errors.get("连号发票", []) or []:
        reason = entry.get("问题原因", "")
        if "需要额外添加支付说明与支付记录" not in reason:
            continue
        items = entry.get("所有重复发票", [])
        inv_objs = [invoices_by_source[item["文件名"]] for item in items if item["文件名"] in invoices_by_source]
        if inv_objs:
            candidates.append(inv_objs)

    grouped_files = {inv["文件名"] for group in candidates for inv in group}
    for category, entries in errors.items():
        if category == "连号发票":
            continue
        for entry in entries or []:
            reason = str(entry.get("问题原因") or "")
            if "支付说明" not in reason or "支付记录" not in reason:
                continue
            inv = invoices_by_source.get(entry.get("文件名"))
            if inv and inv["文件名"] not in grouped_files:
                candidates.append([inv])

    for inv_objs in candidates:
        inv_objs = [inv for inv in inv_objs if not is_high_unit_price(inv)]
        if not inv_objs:
            continue
        indexes: list[int] = []
        images: list[Path] = []
        missing: list[str] = []
        for inv in inv_objs:
            indexes.append(display_index(inv))
            rec_entry = mapping.get(invoice_key(str(inv.get("文件名") or "")), {})
            invoice_payment_images = image_paths(root, list(rec_entry.get("支付记录", []) or []))
            if not invoice_payment_images:
                missing.append(str(inv.get("文件名") or ""))
            images.extend(invoice_payment_images)

        if missing:
            raise RuntimeError(f"missing payment-record screenshots: {', '.join(missing)}")
        idx_text = index_text(indexes)
        groups.append({
            "title": f"xxx {idx_text} 支付记录",
            "images": sorted(images),
            "output": root / INTERNAL_DIR / "支付记录" / f"xxx_{idx_text}_支付记录.docx",
        })
    return groups


def set_font(run, size: int, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def generate_one(title: str, images: list[Path], output: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(title), size=24, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for image in images:
        run = p2.add_run()
        run.add_picture(str(image), width=Cm(DEFAULT_WIDTH_CM))

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"wrote={output} images={len(images)}")


def generate(args: argparse.Namespace) -> None:
    root = args.root.resolve()

    # Manual mode: --images provided
    if args.images:
        images = collect_images(args.images, args.images_dir, args.pattern, root)
        if not images:
            raise RuntimeError("no images matched")
        output = resolve_path(root, args.output)
        generate_one(args.title, images, output)
        return

    # Auto mode: discover groups from invoice_errors.json
    errors_path = resolve_path(root, args.errors)
    results_path = resolve_path(root, args.results)
    match_record = resolve_path(root, args.match_record)

    if not errors_path.exists():
        raise RuntimeError(f"errors file not found: {errors_path}")

    if not match_record.exists():
        raise RuntimeError(f"match record not found: {match_record}")
    groups = auto_collect_groups_from_record(errors_path, results_path, match_record, root)
    if not groups:
        print("no payment-record groups found in invoice_errors.json")
        return

    for group in groups:
        generate_one(group["title"], group["images"], group["output"])


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    add_root_arg(parser)
    parser.add_argument("--title", default=DEFAULT_TITLE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--width-cm", type=float, default=DEFAULT_WIDTH_CM)
    parser.add_argument("--images", type=Path, nargs="*", default=[])
    parser.add_argument("--images-dir", type=Path, default=None)
    parser.add_argument("--pattern", default="*_支付记录*.jpg")
    parser.add_argument("--errors", type=Path, default=DEFAULT_ERRORS, help="invoice_errors.json path (auto mode)")
    parser.add_argument("--results", type=Path, default=DEFAULT_RESULTS, help="invoice_results_sorted.json path")
    parser.add_argument("--match-record", type=Path, default=DEFAULT_MATCH_RECORD)
    return parser.parse_args()


if __name__ == "__main__":
    generate(parse_args())
